"""
論文まとめをWordファイルに出力する
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / 'results'

def set_table_borders(table):
    """テーブルに枠線を設定"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def shade_cell(cell, color='D9E1F2'):
    """セルに背景色を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(4)
    return h

def add_para(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_table(doc, headers, rows, col_widths=None, header_color='2E75B6', header_text_color='FFFFFF'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    set_table_borders(table)
    # ヘッダー行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(header_text_color)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, header_color)
    # データ行
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                shade_cell(cell, 'EBF3FB')
    # 列幅
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def build_document():
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # デフォルトフォント
    doc.styles['Normal'].font.name  = 'Yu Gothic'
    doc.styles['Normal'].font.size  = Pt(10.5)
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # ===========================================================
    # タイトル
    # ===========================================================
    title = doc.add_heading('グラフェン／イオン液体界面におけるナトリウムイオン拡散挙動の'
                            'トポロジカル解析と機械学習予測', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)

    subtitle = doc.add_paragraph('研究成果まとめ')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # ===========================================================
    # 1. 研究背景・目的
    # ===========================================================
    add_heading(doc, '1. 研究背景・目的')
    add_para(doc,
        '全固体電池の固体電解質にイオン液体を添加するとグラフェン電極界面のイオン導電性が向上することが知られている。'
        'しかし、界面近傍のイオンの空間的不均一性が輸送特性に与える影響は複雑であり、最適な界面設計指針の確立が求められている。'
        '本研究では、全原子分子動力学（MD）シミュレーションから得られた原子座標データに対して'
        'パーシステントホモロジー（PH）および動径分布関数（RDF）を構造記述子として適用し、'
        'Na⁺自己拡散係数を目的変数とした機械学習モデルを構築・比較することで、'
        '各記述子の予測精度と物理的解釈力を評価した。'
    )

    # ===========================================================
    # 2. 計算モデル・データセット
    # ===========================================================
    add_heading(doc, '2. 計算モデル・データセット')

    add_heading(doc, '2.1 シミュレーション系', level=2)
    add_para(doc, '系：グラフェン電極 / イオン液体（EMI-BF4 または EMI-TFSI）界面')
    sim_data = [
        ['組成数', '各イオン液体 5種類'],
        ['温度', '250, 300, 350, 400, 450 K（5水準）'],
        ['外部電場', 'なし（ef=0）・あり（ef=1）の2条件'],
        ['総シミュレーション数', '100（BF4系50 + TFSI系50）'],
        ['計算時間', '各100 ns（dt = 1 fs、10ファイル×10 ns）'],
        ['ソフトウェア', 'LAMMPS（OpenMPI並列、64コア）'],
        ['力場', 'OPLS（Lopes et al.）、電荷：AM1-BCC法'],
        ['全原子数', '4,269（Na⁺: 79、EMI⁺: 1,215、BF₄⁻/TFSI⁻: 計1,070または1,440、グラフェンC: 960）'],
    ]
    add_table(doc, ['項目', '内容'], sim_data, col_widths=[5, 11])
    doc.add_paragraph()

    add_heading(doc, '2.2 学習データセット', level=2)
    add_para(doc,
        '各シミュレーションから5スナップショット（60, 70, 80, 90, 100 ns時点）を採取し、'
        '合計500サンプルを使用した（インデックス p000–p499）。'
        'データ分割はシミュレーション単位のGroupKFold（100グループ）を採用し、'
        '訓練70%（350サンプル）・テスト30%（150サンプル）に分割してデータリークを完全防止した。'
    )

    # ===========================================================
    # 3. 拡散係数の再計算
    # ===========================================================
    add_heading(doc, '3. 拡散係数の再計算')

    add_heading(doc, '3.1 元の問題', level=2)
    add_para(doc,
        '既存の計算（MSDの10 ns窓線形フィット）では500サンプル中196件（39.2%）が負値となっており、'
        '物理的に無意味であった。低温域における10 nsウィンドウでのS/N比の低下が原因と判定した。'
    )

    add_heading(doc, '3.2 再計算手法', level=2)
    add_para(doc, '累積拡散係数を以下の式で定義し、各セグメントファイルの最終行のMSD値を使用した：')
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq = p_eq.add_run('D_cumul = MSD(T) / (6T)')
    run_eq.italic = True
    run_eq.font.size = Pt(12)

    add_heading(doc, '3.3 結果', level=2)
    comp_data = [
        ['負値件数', '196/500（39.2%）', '0/500（0%）'],
        ['範囲（Å²/ns）', '−22,530〜163,468', '0.10〜14,085'],
        ['中央値（Å²/ns）', '—', '13.71'],
    ]
    add_table(doc, ['指標', '元の計算', '再計算（本研究）'], comp_data, col_widths=[5, 6, 6])
    doc.add_paragraph()

    add_para(doc, '温度別中央値（Å²/ns）：')
    temp_data = [
        ['250 K', '3.29'], ['300 K', '7.89'], ['350 K', '13.57'],
        ['400 K', '33.57'], ['450 K', '58.33'],
    ]
    add_table(doc, ['温度', '中央値 D_cumul（Å²/ns）'], temp_data, col_widths=[5, 7])
    doc.add_paragraph()
    add_para(doc, '目的変数はlog1p(D_cumul)に変換して使用した（値域：0.09〜9.55）。')

    # ===========================================================
    # 4. 構造記述子
    # ===========================================================
    add_heading(doc, '4. 構造記述子')

    add_heading(doc, '4.1 パーシステントホモロジーベクトル（pd_vector）', level=2)
    add_para(doc,
        'HomCloudライブラリを用いて全4,269原子の座標からアルファ複体フィルトレーションを計算し、'
        '2次のパーシステント図（H2図：3次元空洞に対応）をベクトル化した。'
    )
    pd_vec_data = [
        ['ベクトル化パラメータ', 'PIVectorizerMesh((0, 20), 64, sigma=0.002, weight=("atan", 0.01, 3))'],
        ['フィルトレーション値', 'α = r²（単位: Å²）'],
        ['次元数', '64×64グリッドの上三角部分 = 2,080次元'],
        ['重み関数', 'arctan(0.01 × max(p − 3, 0))　※永続度 p < 3 はweight = 0'],
        ['有効ビン数', '672次元（VarianceThreshold閾値10⁻¹⁰後）'],
    ]
    add_table(doc, ['項目', '内容'], pd_vec_data, col_widths=[5, 11])
    doc.add_paragraph()

    add_heading(doc, '4.2 動径分布関数（RDF）', level=2)
    add_para(doc,
        'LAMMPSのcompute rdfコマンドで各10 nsセグメントにわたり計算済みのRDFファイルを利用した。'
        'ペア：Na⁺-全原子（Na-all）、600ビン（0〜12 Å）。VarianceThreshold後503次元。'
    )

    add_heading(doc, '4.3 Na⁺ z密度プロファイル', level=2)
    add_para(doc,
        'z方向数密度プロファイル（50ビン、換算座標 0〜1）をLAMMPS出力から取得。'
        'グラフェン電極（z ≈ −3.35 Å、換算座標≈0.57）を基準としたNa⁺の界面垂直方向分布を表す。'
    )

    add_heading(doc, '4.4 新PH界面特化記述子（本研究で新規提案）', level=2)
    add_para(doc,
        '全500サンプルのhomcloudテキストファイル（xyz座標）から、境界マップなしのアルファ複体計算'
        '（6コア並列、4.6分）により最大永続性H2特徴量を抽出した。'
    )
    new_ph_data = [
        ['birth α (Å²)', '最大永続性H2特徴量の誕生値', '+0.290***'],
        ['death α (Å²)', '同 消滅値', '+0.296***'],
        ['persistence (Å²)', '同 永続度（death − birth）', '+0.334***'],
        ['n_pairs_gt3', '永続度 > 3 のH2ペア数', '+0.064（n.s.）'],
        ['Na近傍密度', 'z密度換算座標0.55–0.70のNa⁺数密度和（グラフェン近傍）', '+0.160***'],
        ['Na端部密度', 'z密度換算座標0.00–0.19のNa⁺数密度和（ボックス端部）', '+0.592***'],
    ]
    add_table(doc, ['記述子', '内容', 'D との Pearson r'],
              new_ph_data, col_widths=[4, 9, 4])
    doc.add_paragraph()

    add_para(doc, '統計量（500サンプル）：')
    stats_data = [
        ['birth α', '3.39〜880.3 Å²', '中央値 4.90 Å²',
         'BF4: 3.39〜8.80（中央値4.51）、TFSI: 4.16〜880.3（中央値5.39）'],
        ['persistence', '3.53〜63.7 Å²', '中央値 5.85 Å²', '—'],
        ['n_pairs_gt3', '5〜36個', '中央値 16個', '—'],
    ]
    add_table(doc, ['記述子', '範囲', '中央値', '備考'],
              stats_data, col_widths=[3.5, 4, 3, 6])
    doc.add_paragraph()

    # ===========================================================
    # 5. 機械学習モデル
    # ===========================================================
    add_heading(doc, '5. 機械学習モデル')

    add_heading(doc, '5.1 モデル設定', level=2)
    model_settings = [
        ['アルゴリズム', 'ExtraTrees回帰（n_estimators=500, max_depth=5, max_features=0.5, min_samples_leaf=2）'],
        ['特徴量前処理', 'VarianceThreshold → StandardScaler → PCA（20成分）'],
        ['目的変数', 'log1p(D_cumul)'],
        ['データ分割', 'GroupKFold（シミュレーション単位，100グループ）：訓練70%・テスト30%'],
        ['主評価指標', 'GroupKFold-5 クロスバリデーション R²（CV R²）'],
        ['補助指標', 'Test R²、RMSE（実スケール、Å²/ns）'],
    ]
    add_table(doc, ['項目', '設定'], model_settings, col_widths=[5, 11])
    doc.add_paragraph()

    add_heading(doc, '5.2 比較モデルと結果', level=2)
    ml_data = [
        ['RDF + z密度 + 新PH + 物理', '0.735', '0.605', '580', '最高CV R²'],
        ['RDF + z密度 + 物理（最良）', '0.727', '0.617', '334', '最高Test R²・最低RMSE'],
        ['pd_vector + 新PH + 物理', '0.687', '0.592', '638', '—'],
        ['新PH記述子のみ + 物理', '0.682', '0.574', '662', '5次元で pd_vector（0.643）超え'],
        ['RDF のみ + 物理', '0.619', '0.556', '736', '—'],
        ['z密度のみ + 物理', '0.564', '0.487', '602', '—'],
        ['pd_vector + 物理（v3参照）', '0.643', '0.592', '841', '—'],
        ['pd_vectorのみ', '0.420', '0.361', '944', '—'],
    ]
    add_table(doc,
              ['モデル', 'CV R²', 'Test R²', 'RMSE（Å²/ns）', '備考'],
              ml_data, col_widths=[6.5, 1.8, 1.8, 3.5, 4])
    doc.add_paragraph()
    add_para(doc, '※物理パラメータ：IL種（BF4/TFSI）、温度、組成、電場、スナップショット時刻の5次元。')

    add_heading(doc, '5.3 特徴量重要度（RDF + z密度 + 新PH + 物理モデル）', level=2)
    imp_data = [
        ['z密度 PCA', '34.6%'],
        ['物理パラメータ', '24.5%'],
        ['新PH記述子（計）', '23.0%'],
        ['RDF PCA', '17.9%'],
    ]
    add_table(doc, ['特徴量グループ', '重要度'], imp_data, col_widths=[8, 4])
    doc.add_paragraph()
    add_para(doc,
        '新PH記述子23.0%の内訳：'
        'Na端部密度 18.5%、Na近傍密度 3.7%、birth/death/persistence/n_pairs_gt3 合計 0.8%。'
    )

    # ===========================================================
    # 6. パーシステントホモロジー逆解析
    # ===========================================================
    add_heading(doc, '6. パーシステントホモロジー逆解析')

    add_heading(doc, '6.1 手法', level=2)
    add_para(doc,
        'HomCloud v5.3の optimal_volume()（境界マップあり）を用いて、'
        '最大永続性H2特徴量の境界原子群を特定した。'
        '対象：BF4_gra No.01 の100 nsスナップショット（250 K・450 K）。'
    )

    add_heading(doc, '6.2 結果', level=2)
    inv_data = [
        ['最大永続性H2ペア\nbirth / death / persistence', '4.37 / 8.70 / 4.33 Å²', '4.89 / 13.29 / 8.40 Å²'],
        ['永続度 > 3 のペア数（全H2ペア数）', '11個（6,101個中）', '14個（5,366個中）'],
        ['境界原子数（合計）', '155個', '52個'],
        ['C（グラフェン）', '43個（27.7%）', '0個（0%）'],
        ['F（BF₄⁻）', '49個（31.6%）', '15個（28.8%）'],
        ['Na⁺', '6個（3.9%）', '1個（1.9%）'],
        ['EMI⁺（C/N/H）', '57個（36.8%）', '36個（69.2%）'],
        ['境界原子z範囲', '−3.35〜+11.26 Å\n（グラフェン面から11.26 Å上方）', '−0.56〜+8.26 Å\n（グラフェン面と非接触）'],
        ['物理的解釈', 'グラフェン43原子が底蓋を形成し\n6個のNa⁺を三次元封じ込め\n→ 界面空洞構造', 'グラフェン非関与\nバルクIL内部の空洞'],
    ]
    add_table(doc, ['項目', '250 K（低D系）', '450 K（高D系）'],
              inv_data, col_widths=[5.5, 6, 6])
    doc.add_paragraph()

    add_heading(doc, '6.3 RDF・z密度との定量的関係（N = 500）', level=2)
    corr_data = [
        ['RDF各ビン vs log1p(D)', 'r = +0.45（r = 4.7 Å, 第3配位殻）'],
        ['z密度各ビン vs log1p(D)', 'r = +0.63（z換算座標 = 0.99）'],
        ['pd_vector各ビン vs log1p(D)', 'r = −0.56（birth_grid ≈ 0.47 Å²）'],
        ['RDF PC vs pd_vector PC', '最大 r = 0.78（情報重複）'],
        ['z密度 PC vs pd_vector PC', '最大 r = 0.71（情報重複）'],
        ['RDF PC vs z密度 PC', '最大 r = 0.94（ほぼ同一情報）'],
        ['pd_vector残差解析', 'RDF除去後の残差へのpd_vector寄与：CV R² = −0.11'],
    ]
    add_table(doc, ['解析', '結果'], corr_data, col_widths=[7, 9.5])
    doc.add_paragraph()

    # ===========================================================
    # 7. 考察
    # ===========================================================
    add_heading(doc, '7. 考察')

    add_heading(doc, '7.1 RDFが予測精度で優位な理由', level=2)
    reasons = [
        '情報の直接性：RDFのNa-all第2・第3配位殻ピーク（r = 3.65〜5.39 Å）は拡散係数と直接相関（最大 r = +0.45）。配位殻の緩さ（エネルギー障壁の低さ）を直接反映する。',
        '情報の重複：pd_vectorとRDFは互いに高相関（最大 r = 0.78）であり、pd_vectorが捉える情報の多くはRDFで説明済み（偏相関：pd_vector PC1のDとの直接相関 r = −0.37 → RDF除去後 r = −0.09）。',
    ]
    for r in reasons:
        p = doc.add_paragraph(style='List Number')
        p.add_run(r).font.size = Pt(10.5)

    add_heading(doc, '7.2 pd_vectorの予測精度が低い理由', level=2)
    reasons2 = [
        '稀少性問題：界面空洞（物理的に最重要なH2特徴量）は6,101ペア中のたった1個であり、短命なペアが2,080次元pd_vectorを支配する。',
        '重み関数の制約：arctan(0.01 × max(p−3, 0))により永続度 < 3 のペアはweight = 0。ゼロ重みペアのGaussianにじみが支配的になる。',
        'スケールの非一致：pd_vectorの有効範囲（0〜20 Å²、r ≤ 4.47 Å）が最大永続性H2特徴量のbirth（4.37〜4.89 Å²）の境界付近に対応し、高精度な捕捉が困難。',
    ]
    for r in reasons2:
        p = doc.add_paragraph(style='List Number')
        p.add_run(r).font.size = Pt(10.5)

    add_heading(doc, '7.3 新PH記述子の有効性と限界', level=2)
    add_para(doc,
        '【有効性】「新PH+物理」（CV R² = 0.682）は「pd_vector+物理」（CV R² = 0.643）を上回り、'
        '最大永続性H2特徴量の5次元圧縮が2,080次元PCA-20より効率的であることが示された。'
    )
    add_para(doc,
        '【限界】真に重要だったのはbirth/death/persistence値（重要度合計0.8%）ではなく、'
        'Na端部密度（18.5%）というz密度プロキシであった。birth値はBF4（4.51 Å²）と'
        'TFSI（5.39 Å²、最大880 Å²）の違いを反映するが、IL種は物理パラメータで既に捕捉されている。'
        'また、TFSIではbirth最大880 Å²の外れ値が存在し、RDFモデルへの追加でTest R²が低下（0.617→0.605）、'
        'RMSEが悪化（334→580 Å²/ns）する。'
    )

    add_heading(doc, '7.4 PHの独自の価値（逆解析による物理的解釈）', level=2)
    add_para(doc,
        'PHの予測精度面での優位性は示されなかったが、逆解析による原子論的可視化に不可欠な情報を提供する。'
        '250 KのNa⁺閉じ込めを「グラフェンC43個 + BF₄⁻49個 + Na⁺6個」の三次元空洞として定量的に同定し、'
        '450 Kではグラフェン非関与のバルクIL空洞が支配的であることを直接実証した。'
        'これはRDFやz密度（「Na⁺がどこにいるか」）では不可能な、'
        '「Na⁺がどのような三次元かご構造に封じ込められているか」という高次構造情報であり、'
        'RDF・z密度との間に大きな情報重複（最大 r = 0.94）があるにもかかわらず、'
        'PHのみが与えられる本質的知見である。'
    )

    # ===========================================================
    # 8. 結論
    # ===========================================================
    add_heading(doc, '8. 結論')
    conclusions = [
        '拡散係数の再計算：累積拡散係数 D_cumul = MSD(T)/(6T) を用いることで、従来手法で196件（39.2%）存在した負値を完全に排除し（D = 0.10〜14,085 Å²/ns）、500サンプル全てで正値の学習データを構築した。',
        '最良予測モデル：RDF（Na-all g(r)）+ Na⁺ z密度プロファイル + 物理パラメータを組み合わせたExtraTreesモデルが最高の汎化精度（CV R² = 0.727、Test R² = 0.617、RMSE = 334 Å²/ns）を達成した。',
        '記述子の予測力比較：RDF > z密度 > 新PH界面特化記述子 > pd_vector（2,080次元）の順。PHが5次元の界面特化記述子（CV R² = 0.682）として2,080次元pd_vector（CV R² = 0.643）を上回ることを実証した。',
        'PHの優位性の所在：予測精度ではなく界面構造の解釈にある。逆解析により、低温（低D）ではグラフェンC43個が底蓋を形成しNa⁺6個を三次元空洞に封じ込める界面トポロジー構造が存在し、高温（高D）ではこの構造が消失してバルクIL空洞に置換されることを原子レベルで実証した。このメカニズムはRDFおよびz密度では記述不可能な高次トポロジー情報である。',
        '予測精度の限界と展望：現在のCV R² ≈ 0.73の主要因はサンプル数（100シミュレーション）の制限と記述子の情報重複にある。PHを活用した改善の余地として、境界マップ付き逆解析による「グラフェン関与割合」・「封じ込めNa⁺数」の特化記述子が有望である。',
    ]
    for c in conclusions:
        p = doc.add_paragraph(style='List Number')
        p.add_run(c).font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(4)

    # ===========================================================
    # 9. 生成ファイル一覧
    # ===========================================================
    add_heading(doc, '9. 生成ファイル一覧')
    files_data = [
        ['D_cumulative.csv', '再計算拡散係数（500サンプル）'],
        ['diffusion_ml_v3.py', '主要MLスクリプト（v3最終版）'],
        ['ml_results_v3.csv/png', '主要MLモデル比較（7モデル）'],
        ['rdf_ml_results.csv', 'RDF・z密度モデル比較（7モデル）'],
        ['ph_targeted_results.csv', '新PH記述子を含むモデル比較（7モデル）'],
        ['pd_corr_map.png', 'pd_vectorビン vs D相関マップ（2D birth-death空間）'],
        ['pd_residual_analysis.png', '残差解析結果（pd_vectorの付加情報量）'],
        ['pd_h2_diagrams.png', '条件別H2ダイアグラム（温度・IL種・電場）'],
        ['rdf_zdensity_analysis.png', 'RDF・z密度 vs D相関・ML比較'],
        ['homcloud_pd_comparison.png', '250 K・450 K H2パーシステンス図比較'],
        ['homcloud_inverse_3d.png', '逆解析結果3D可視化'],
        ['homcloud_xy_projection.png', '逆解析境界原子のグラフェン面投影'],
        ['homcloud_comprehensive.png', '総合可視化（温度・条件別・定量トレンド）'],
        ['ph_targeted_ml.png', '新PH記述子ML比較・特徴量重要度'],
    ]
    add_table(doc, ['ファイル名', '内容'], files_data, col_widths=[6.5, 10])

    # ===========================================================
    # 保存
    # ===========================================================
    out_path = str(OUT_DIR / '研究成果まとめ.docx')
    doc.save(out_path)
    print(f'保存完了: {out_path}')
    return out_path

if __name__ == '__main__':
    build_document()
